#!/usr/bin/env python3
"""
Word/DOCX Tool — 读取、创建、转换 Word 文档。

Commands: read, create, convert, template
"""

import argparse
import json
import os
import sys
from datetime import datetime
from pathlib import Path

OUTPUT_DIR = Path.home() / ".openclaw" / "output" / "word"


def _now():
    return datetime.now().isoformat(timespec="seconds")


def _ensure_dir():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return OUTPUT_DIR


def cmd_read(filepath: str) -> dict:
    """Read a Word document and output text + structure."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    if not os.path.isfile(filepath):
        return {"error": f"File not found: {filepath}"}

    try:
        doc = Document(filepath)
        paragraphs = []
        for p in doc.paragraphs:
            paragraphs.append({
                "text": p.text,
                "style": p.style.name if p.style else "Normal",
            })

        tables = []
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append({"index": i, "rows": len(rows), "data": rows[:5]})

        result = {
            "file": filepath,
            "paragraphs": len(paragraphs),
            "tables": len(tables),
            "preview": [p["text"] for p in paragraphs[:20] if p["text"].strip()],
            "structure": paragraphs[:50],
            "table_preview": tables[:3],
        }

        print(f"✅ 读取成功: {filepath} ({len(paragraphs)}段, {len(tables)}表)")
        return result

    except Exception as e:
        return {"error": str(e)}


def cmd_create(title: str, content: str, output: str = None) -> dict:
    """Create a Word document from title + markdown-ish content."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    _ensure_dir()
    if not output:
        safe = title.replace(" ", "_").replace("/", "_")[:30]
        output = str(OUTPUT_DIR / f"{safe}_{datetime.now().strftime('%Y%m%d')}.docx")

    try:
        doc = Document()
        doc.add_heading(title, level=0)

        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("1. ") or line.startswith("2. ") or line.startswith("3. "):
                doc.add_paragraph(line[3:], style="List Number")
            else:
                doc.add_paragraph(line)

        doc.save(output)
        print(f"✅ 创建成功: {output}")
        return {"file": output, "title": title, "timestamp": _now()}

    except Exception as e:
        return {"error": str(e)}


def cmd_convert(filepath: str, to_format: str = "md") -> dict:
    """Convert Word to Markdown or plain text."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed"}

    if not os.path.isfile(filepath):
        return {"error": f"File not found: {filepath}"}

    _ensure_dir()
    base = Path(filepath).stem

    try:
        doc = Document(filepath)
        lines = []

        for p in doc.paragraphs:
            style = p.style.name if p.style else "Normal"
            text = p.text.strip()
            if not text:
                lines.append("")
                continue

            if to_format == "md":
                if "Heading 1" in style:
                    lines.append(f"# {text}")
                elif "Heading 2" in style:
                    lines.append(f"## {text}")
                elif "Heading 3" in style:
                    lines.append(f"### {text}")
                elif "List Bullet" in style:
                    lines.append(f"- {text}")
                elif "List Number" in style:
                    lines.append(f"1. {text}")
                else:
                    lines.append(text)
            else:
                lines.append(text)

        ext = "md" if to_format == "md" else "txt"
        out = str(OUTPUT_DIR / f"{base}.{ext}")
        with open(out, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))

        print(f"✅ 转换完成: {out}")
        return {"output": out, "format": to_format, "paragraphs": len(lines)}

    except Exception as e:
        return {"error": str(e)}


def health_check() -> dict:
    checks = [{"name": "python3", "status": "ok"}]
    for mod, pkg in [("docx", "python-docx")]:
        try:
            __import__(mod)
            checks.append({"name": pkg, "status": "ok"})
        except ImportError:
            checks.append({"name": pkg, "status": "warn", "message": f"pip install {pkg}"})
    overall = "warn" if any(c["status"] == "warn" for c in checks) else "ok"
    return {"skill": "word-docx", "version": "1.0.0", "status": overall,
            "checks": checks, "timestamp": _now()}


def main():
    parser = argparse.ArgumentParser(description="Word/DOCX Tool")
    parser.add_argument("--check", action="store_true")
    parser.add_argument("--json", action="store_true")

    sub = parser.add_subparsers(dest="command")

    p_read = sub.add_parser("read")
    p_read.add_argument("file")

    p_create = sub.add_parser("create")
    p_create.add_argument("--title", required=True)
    p_create.add_argument("--content", required=True)
    p_create.add_argument("--output")

    p_convert = sub.add_parser("convert")
    p_convert.add_argument("file")
    p_convert.add_argument("--format", default="md", choices=["md", "txt"])

    args = parser.parse_args()

    if args.check:
        r = health_check()
        print(json.dumps(r, indent=2, ensure_ascii=False))
        sys.exit(0 if r["status"] != "fail" else 1)

    if not args.command:
        parser.print_help()
        sys.exit(0)

    result = None
    if args.command == "read":
        result = cmd_read(args.file)
    elif args.command == "create":
        result = cmd_create(args.title, args.content, args.output)
    elif args.command == "convert":
        result = cmd_convert(args.file, args.format)

    if args.json and result:
        print(json.dumps(result, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
